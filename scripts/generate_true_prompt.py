"""Генерация true_prompt.docx — промпт для воссоздания идентичного KeepCopy."""

from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "true_prompt.docx"


def _shade_cell(cell, fill: str) -> None:
    shading = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    shd.set(qn("w:val"), "clear")
    shading.append(shd)


def _add_title_page(doc: Document) -> None:
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Промпт / техническое задание\n")
    run.bold = True
    run.font.size = Pt(22)
    run2 = title.add_run("KeepCopy")
    run2.bold = True
    run2.font.size = Pt(28)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Полная спецификация для создания идентичного приложения с нуля"
    )
    r.font.size = Pt(13)
    r.italic = True

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(
        f"Версия документа: 2.0\nДата: {date.today().strftime('%d.%m.%Y')}\n"
        "Источник: актуальный код проекта KeepCopy"
    )

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nr = note.add_run(
        "Используй этот документ как единственный источник требований. "
        "Реализуй поведение один-в-один: UI, логику копирования, расписание, "
        "автозапуск, хранение данных, сборку и тесты. Не упрощай и не опускай "
        "описанные UX-детали."
    )
    nr.font.size = Pt(10)
    nr.italic = True
    doc.add_page_break()


def _add_info_box(doc: Document, title: str, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.rows[0].cells[0]
    _shade_cell(cell, "E8F0FE")
    p = cell.paragraphs[0]
    r = p.add_run(title + "\n")
    r.bold = True
    for line in lines:
        cell.add_paragraph(line, style="List Bullet")
    doc.add_paragraph()


def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
        _shade_cell(hdr[i], "F2F2F2")
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = text
    doc.add_paragraph()


def _bullets(doc: Document, items: list[str], style: str = "List Bullet") -> None:
    for item in items:
        doc.add_paragraph(item, style=style)


def build_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    _add_title_page(doc)

    # --- 0. Инструкция исполнителю ---
    doc.add_heading("0. Инструкция исполнителю (промпт)", level=1)
    doc.add_paragraph(
        "Создай настольное Windows-приложение «KeepCopy» строго по этому ТЗ. "
        "Язык кода — Python 3.10+, UI — PySide6, фильтры — pathspec (gitwildmatch), "
        "хранение — JSON, сборка — Nuitka standalone. Интерфейс и сообщения — на русском."
    )
    _bullets(
        doc,
        [
            "Сначала каркас: models → services → workers → ui → main.py → tests → scripts сборки.",
            "Покрой pytest-тестами логику копирования, пути, автозапуск, кэш размеров, валидацию задачи.",
            "Не добавляй лишних фич вне ТЗ; не меняй семантику режимов копирования и расписания.",
            "Особо внимательно реализуй UX таблицы: раскрытие источников, размеры, выделение, фон кнопок действий.",
        ],
    )

    # --- 1. Общие сведения ---
    doc.add_heading("1. Общие сведения", level=1)
    doc.add_paragraph(
        "KeepCopy — приложение для ручного и периодического резервного копирования "
        "файлов и каталогов. Пользователь задаёт источники, папку назначения, "
        "расписание, исключения и режим копирования. Копирование и подсчёт размеров "
        "идут в фоне без блокировки UI."
    )
    _add_table(
        doc,
        ["Параметр", "Значение"],
        [
            ["Название", "KeepCopy"],
            ["Целевая ОС", "Windows 10 и новее"],
            ["Язык UI", "Русский"],
            ["Точка входа", "main.py / KeepCopy.exe"],
            ["Данные", "settings.json рядом с программой"],
            ["Логи", "backup.log (ротация 1000 строк, UTF-8)"],
            ["Ошибки копирования", "папка errors/"],
            ["Зависимости runtime", "PySide6, pathspec"],
            ["Тесты", "pytest"],
            ["Сборка", "Nuitka → compiler/KeepCopy.dist/"],
        ],
    )

    # --- 2. Глоссарий ---
    doc.add_heading("2. Глоссарий", level=1)
    for term, definition in [
        ("Задача", "Настроенное копирование: источники, назначение, расписание, режим, сжатие."),
        ("Источник", "Файл или папка; в задаче может быть несколько."),
        ("Назначение", "Существующая папка назначения (с буквой диска)."),
        ("Исключение", "Шаблон gitwildmatch относительно корня источника."),
        ("Слой", "Папка или ZIP backup_ДД.ММ.ГГГГ_NNN в режимах «слоями» и «дублирование»."),
        ("Активная задача", "is_active=true; участвует в авторасписании."),
        ("След.вып.", "Дата следующего автозапуска (без времени в ячейке); меняется только после автозапуска."),
        ("Послед.вып.", "Дата и время последнего запуска (ручного или авто)."),
    ]:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"{term} — ")
        r.bold = True
        p.add_run(definition)

    # --- 3. Запуск и жизненный цикл ---
    doc.add_heading("3. Запуск, трей и закрытие", level=1)
    doc.add_heading("3.1. Точка входа", level=2)
    _bullets(
        doc,
        [
            "Аргумент --background: старт со скрытым окном (автозапуск Windows).",
            "При старте: load settings.json → reconcile_autostart → MainWindow.",
            "stay_in_background = (--background) ИЛИ settings.autostart.",
            "Если stay_in_background: setQuitOnLastWindowClosed(False), создать AppTray.",
            "При --background: warmup окна (show off-screen → hide), чтобы рамка HWND была корректной при первом показе из трея.",
            "Без --background, но с автозапуском: окно показывается сразу, трей уже есть.",
            "Без автозапуска и без --background: обычное окно, трей не нужен.",
        ],
    )

    doc.add_heading("3.2. Трей", level=2)
    _bullets(
        doc,
        [
            "Иконка «KeepCopy»; меню: «Открыть», «Выход».",
            "Двойной клик / «Открыть» — показать и активировать окно.",
            "«Выход» — полное завершение (_force_quit, остановить планировщик, quit).",
        ],
    )

    doc.add_heading("3.3. Закрытие окна (×)", level=2)
    _add_info_box(
        doc,
        "Правило закрытия",
        [
            "Если автозапуск включён ИЛИ процесс стартовал с --background → скрыть в трей (не завершать).",
            "Иначе → остановить планировщик и закрыть окно (процесс завершается).",
            "При включении автозапуска в настройках текущей сессии — сразу создать трей и переключить режим закрытия.",
            "При выключении автозапуска (и если не было --background) — убрать трей; закрытие снова завершает приложение.",
        ],
    )

    # --- 4. UI ---
    doc.add_heading("4. Пользовательский интерфейс", level=1)
    doc.add_paragraph(
        "PySide6, стиль Fusion, темы light/dark (бирюзовая палитра COLORS + QSS). "
        "SVG-иконки 16×16. Курсор «рука» на интерактивных элементах. "
        "На Windows 11 — цвет заголовка через DWM; на Windows 10 — immersive dark/light."
    )

    doc.add_heading("4.1. Главное окно", level=2)
    doc.add_paragraph("Заголовок «KeepCopy». Панель инструментов | таблица задач | строка статуса.")
    _add_table(
        doc,
        ["Кнопка", "Поведение"],
        [
            ["Создать", "Диалог новой задачи"],
            [
                "Выполнить",
                "Выделенные задачи; если нет выделения — все активные; если активных нет — все. "
                "Блокируется, пока занята очередь",
            ],
            ["Настройки", "Тема, колонки, шрифт, автозапуск"],
            ["Помощь", "Краткая HTML-справка"],
        ],
    )

    doc.add_heading("4.2. Таблица задач — колонки", level=2)
    doc.add_paragraph(
        "Порядок фиксирован. Колонка «Действия» всегда видима. Перед колонками — служебный № (ширина 30)."
    )
    _add_table(
        doc,
        ["Ключ", "Заголовок", "По умолч.", "Ширина", "Примечание"],
        [
            ["name", "Название", "да", "150", ""],
            ["description", "Описание", "нет", "120", ""],
            ["sources", "Источники", "да", "305", "первый путь + « +N»; elide"],
            ["destination", "Назначение", "да", "160", "elide; открытие в проводнике"],
            ["schedule_time", "Время", "да", "60", "ЧЧ:ММ"],
            ["period_days", "Период.", "да", "65", "дни"],
            ["exclusions", "Исключения", "нет", "120", "первый + « +N»"],
            ["total_size", "Размер", "да", "80", "см. §4.4"],
            ["last_run", "Послед.вып.", "да", "115", "ДД.ММ.ГГГГ ЧЧ:ММ"],
            ["next_run", "След.вып.", "да", "95", "только дата; у неактивных — приглушённый"],
            ["copy_mode", "Режим копирования", "да", "145", "метка режима"],
            ["compress", "Сжатие", "да", "65", "«да»/«нет»"],
            ["actions", "Действия", "всегда", "90", "play/pause, edit, delete"],
        ],
    )
    doc.add_paragraph(
        "Ширина окна = сумма видимых колонок (шаг округления 5). "
        "Высота = max(600 px, 60% высоты экрана). Ширины колонок сохраняются в settings.json."
    )

    doc.add_heading("4.3. Выделение, hover, кнопки действий", level=2)
    _bullets(
        doc,
        [
            "Своё выделение строк (не Qt SelectRows): клик — выбрать; Ctrl — toggle; Shift — диапазон по строкам задач.",
            "Строки деталей не выделяются. Клик вне таблицы / по пустому месту — сброс выделения.",
            "Qt-режим NoSelection; сбрасывать «текущую ячейку», чтобы не было системной подсветки одной клетки.",
            "Hover строки: полупрозрачный цвет выделения (alpha 128); select — непрозрачный (alpha 255).",
            "Неактивные задачи — приглушённый цвет текста (#8aa4ad light).",
            "Колонка «Действия»: контейнер ActionCellContainer рисует фон строки сам (base + highlight), "
            "чтобы кнопки не имели «чужого» фона. Учитывать чередование цветов строк (Base / AlternateBase).",
            "Иконки: play = задача активна («Приостановить»); pause = неактивна («Активировать»); "
            "delete — красный hover, белый крестик.",
            "Не менять QSS контейнера при каждом выделении — иначе шрифт таблицы визуально «худеет».",
        ],
    )

    doc.add_heading("4.4. Раскрытие источников / исключений и размеры", level=2)
    _add_info_box(
        doc,
        "Критичные UX-правила размеров и списка",
        [
            "Клик по ячейке «Источники» или «Исключения» строки задачи раскрывает/сворачивает список, "
            "только если источников > 1 или исключений > 1.",
            "Анимации раскрытия нет: полный refresh таблицы. (Старая анимация defaultSectionSize ломала высоты строк.)",
            "В строке задачи всегда отображается только первый источник/исключение + « +N».",
            "В выпадающем списке — только остальные: sources[1:], exclusions[1:] (первый источник в списке НЕ дублировать).",
            "Текст деталей: «  → путь», цвет detail (#5a7a84 / #9ec4cc).",
            "Кэш размеров по каждому источнику отдельно; в свёрнутой строке задачи — СУММА, обычный цвет шрифта.",
            "В раскрытой строке задачи колонка «Размер» = размер ТОЛЬКО первого источника, цветом detail (как у списка).",
            "У строк дополнительных источников в списке — свой размер (цвет detail). У строк исключений размер пустой.",
            "Пока размер не пересчитан / устарел — «—» или оранжевый #f57c00 до завершения скана.",
        ],
    )

    doc.add_heading("4.5. Прочее поведение таблицы", level=2)
    _bullets(
        doc,
        [
            "Двойной клик / ПКМ по «Источники» или «Назначение» — открыть в проводнике (несколько источников — меню).",
            "Горизонтальный скролл: Alt + колёсико.",
            "Подсказка строки: HTML с названием, назначением (акцент) и всеми источниками.",
            "Статус: «Задача X: подсчёт/копирование — …путь… (i из n) — p%».",
        ],
    )

    doc.add_heading("4.6. Диалог задачи", level=2)
    doc.add_paragraph("Секции: Описание (развёрнута), Исключения (свёрнута), Дополнительно (свёрнута). «Сохранить» по центру.")
    _add_table(
        doc,
        ["Поле", "Требования"],
        [
            ["Название", "Обязательно; по умолчанию «Задача N»"],
            ["Описание", "Необязательно"],
            ["Источник", "≥1; папка или файл; таблица с удалением; вложенные пути красным и НЕ сохраняются"],
            ["Назначение", "Существующая папка с буквой диска"],
            ["Время", "ЧЧ:ММ или пусто (без расписания)"],
            ["Период", "1–366; по умолчанию 7; пусто/0 = без периода"],
            ["Исключения", "gitwildmatch; без дубликатов"],
            ["Макс. размер МБ", "Файлы строго больше лимита пропускаются; равные — копируются"],
            ["Сжатие", "ZIP_DEFLATED, compresslevel=5"],
            ["Режим", "Сохранять изменения / Сохранять слоями / Дублирование"],
        ],
    )
    doc.add_paragraph(
        "Валидация с ошибками под полями. Закрытие с несохранёнными изменениями — предупреждение. "
        "Таблицы источников/исключений: № 30px, действия 30px, путь stretch ≥150; "
        "подсветка строки select/hover; фон кнопки удаления совпадает со строкой."
    )

    doc.add_heading("4.7. Настройки и помощь", level=2)
    _bullets(
        doc,
        [
            "Настройки: тема light/dark, видимость колонок (кроме «Действия»), шрифт 10–16 (default 12), автозапуск.",
            "Только «Применить» сохраняет. «Отмена» — без изменений.",
            "Помощь: краткий HTML по задачам, источникам, исключениям, режимам, таблице, ошибкам.",
        ],
    )

    # --- 5. Модель задачи ---
    doc.add_heading("5. Модель задачи", level=1)
    doc.add_paragraph(
        "Поля: id (uuid), name, description, sources[], destination, schedule_time, period_days, "
        "exclusions[], max_size_mb, compress, copy_mode, is_active, created_at, last_run, "
        "last_auto_run, next_run, errors_counter."
    )
    doc.add_paragraph(
        "При создании по умолчанию: время 00:00, период 7, is_active=true, "
        "next_run = дата_создания + period_days (не в день создания)."
    )

    # --- 6. Копирование ---
    doc.add_heading("6. Логика резервного копирования", level=1)

    doc.add_heading("6.1. Именование в назначении", level=2)
    _bullets(
        doc,
        [
            "Источник-папка → подпапка с именем последнего компонента пути.",
            "Источник-файл → подпапка «имя_родительской_папки_files».",
            "Внутри — относительная структура от корня источника.",
        ],
    )

    doc.add_heading("6.2. Режимы", level=2)
    _add_table(
        doc,
        ["Режим", "Поведение"],
        [
            [
                "Сохранять изменения (keep_changes)",
                "Инкремент по mtime > last_run в одну папку/ZIP. "
                "Ручной запуск также копирует файлы, отсутствующие в назначении. "
                "Первый автозапуск (last_auto_run is None) — все подходящие файлы.",
            ],
            [
                "Сохранять слоями (layered)",
                "Новый слой backup_ДД.ММ.ГГГГ_NNN только с изменениями; пустой слой не создавать.",
            ],
            [
                "Дублирование (duplicate)",
                "Каждый запуск — полный слой (включая пустые каталоги), новый суффикс NNN.",
            ],
        ],
    )

    doc.add_heading("6.3. ZIP и фильтры", level=2)
    _bullets(
        doc,
        [
            "keep_changes + zip: один ZIP на источник (дополнение к существующему).",
            "layered/duplicate + zip: ZIP на каждый слой.",
            "Фильтры — относительно корня источника; совпадение с папкой исключает ветку.",
            "Колонка «Размер» учитывает фильтры и max_size, но не режим и не last_run.",
        ],
    )

    doc.add_heading("6.4. Алгоритм выполнения", level=2)
    for i, step in enumerate(
        [
            "Собрать файлы (FileMatcher) по источникам, фильтрам, max_size и режиму.",
            "Проверить свободное место: нужно total_bytes * 1.10; иначе отмена, лог «недостаточно места», last_run не обновлять (next_run при авто — обновить).",
            "Копировать по источникам; длинные пути Windows — префикс \\\\?\\.",
            "Ошибки отдельных файлов — пропуск; список в errors/errors_ДД.ММ.ГГГГ_NNN; errors_counter++.",
            "Исходники только читать (не удалять/не менять).",
            "Прогресс по байтам, обновление статуса не чаще ~0.35 с (кроме 100%).",
        ],
        1,
    ):
        doc.add_paragraph(step, style="List Number")

    # --- 7. Расписание ---
    doc.add_heading("7. Расписание", level=1)
    _bullets(
        doc,
        [
            "Задача «по расписанию», если заданы и время, и период.",
            "Автозапуск, когда now ≥ след.вып. + время задачи и is_active.",
            "После авто: next = base + ((today-base)//period + 1)*period (догоняющая формула, один прогон).",
            "Ручной запуск не меняет next_run.",
            "Таймер планировщика 30 с; при старте приложения, если есть просроченные — догоняющий запуск через 5 минут.",
            "Пауза (deactivate) сохраняет next_run; при активации — пересчёт.",
        ],
    )

    # --- 8. Очередь ---
    doc.add_heading("8. Очередь и потоки", level=1)
    _bullets(
        doc,
        [
            "SizeScanWorker (QThread): подсчёт → {total, sources{path: bytes}}.",
            "BackupWorker (QThread): BackupEngine.run.",
            "TaskQueueManager: последовательно scan → copy; при старте scan_all_sizes всех задач.",
            "Перед каждым копированием — свежий скан размера этой задачи.",
            "Сигналы размера через object/dict (не только qint64), чтобы передать breakdown по источникам.",
        ],
    )

    # --- 9. Хранение ---
    doc.add_heading("9. Хранение (settings.json)", level=1)
    doc.add_paragraph("Рядом с main.py / KeepCopy.exe. Структура:")
    code = (
        "{\n"
        '  "tasks": [ { /* поля Task */ } ],\n'
        '  "settings": {\n'
        '    "theme": "light"|"dark",\n'
        '    "font_size": 12,\n'
        '    "visible_columns": { "name": true, ... },\n'
        '    "column_widths": { "name": 150, ... },\n'
        '    "autostart": false\n'
        "  },\n"
        '  "size_cache": {\n'
        '    "<task_id>": {\n'
        '      "total": 123456,\n'
        '      "sources": { "C:\\\\path\\\\src": 1000, ... }\n'
        "    }\n"
        "  }\n"
        "}"
    )
    p = doc.add_paragraph()
    r = p.add_run(code)
    r.font.name = "Consolas"
    r.font.size = Pt(9)
    doc.add_paragraph(
        "Обратная совместимость: если size_cache[id] — целое число, трактовать как "
        "{total: N, sources: {}} до следующего скана."
    )

    # --- 10. Логи ---
    doc.add_heading("10. Логирование", level=1)
    _add_table(
        doc,
        ["Файл", "Правила"],
        [
            [
                "backup.log",
                "дата время | задача | источник | описание≤100 | результат. Ротация 1000 строк.",
            ],
            [
                "errors/errors_ДД.ММ.ГГГГ_NNN",
                "Метаданные задачи + список пропущенных файлов. Один файл на запуск с ошибками.",
            ],
        ],
    )
    doc.add_paragraph(
        "Результаты: «успешно»; «скопировано с ошибками (файл …)»; "
        "«ошибка: недостаточно места на диске». На каждый источник — строка лога."
    )

    # --- 11. Автозапуск ---
    doc.add_heading("11. Автозапуск Windows", level=1)
    _add_table(
        doc,
        ["Параметр", "Значение"],
        [
            ["Папка", "%APPDATA%\\Microsoft\\Windows\\Start Menu\\Programs\\Startup"],
            ["Ярлык", "KeepCopy.lnk"],
            ["Цель (exe)", "KeepCopy.exe --background"],
            ["WindowStyle", "7 (свёрнуто)"],
            ["Разработка", "pythonw.exe \"main.py\" --background"],
        ],
    )
    _bullets(
        doc,
        [
            "reconcile_autostart при каждом старте: ярлык ↔ галочка; при включённом — пересоздать ярлык на актуальный путь.",
            "После переноса папки с программой автозапуск сломается, пока не запустят exe из нового места "
            "(тогда ярлык обновится) или не переключат галочку.",
            "Nuitka: определять сборку по sys.argv[0] (.exe, не python) — Nuitka не ставит sys.frozen.",
        ],
    )

    # --- 12. Сборка ---
    doc.add_heading("12. Сборка (Nuitka)", level=1)
    _bullets(
        doc,
        [
            "Скрипт: python scripts/build_exe.py.",
            "NUITKA_CACHE_DIR=C:\\NuitkaCache (короткий путь для MinGW).",
            "standalone, без консоли, plugin pyside6, include pathspec и assets/.",
            "Выход: compiler/KeepCopy.dist/ (переименовать из main.dist); сохранять settings.json при пересборке.",
            "main.build — промежуточная папка; для работы exe не нужна; ускоряет повторные сборки — не удалять скриптом по умолчанию.",
            "Антивирус может ложно блокировать (Bearfoos.A!ml) — исключить всю папку KeepCopy.dist.",
        ],
    )

    # --- 13. Структура проекта ---
    doc.add_heading("13. Структура проекта", level=1)
    _bullets(
        doc,
        [
            "main.py — точка входа",
            "models/ — Task, AppSettings, колонки",
            "services/ — storage, backup_engine, file_matcher, scheduler, autostart, logger, path_utils",
            "workers/ — task_queue, backup_worker, size_scan_worker, auto_scheduler",
            "ui/ — main_window, main_table, task_dialog, settings_dialog, help_dialog, app_tray, "
            "themes, widgets (ActionCellContainer), icons, window_chrome, cursors, message_box, app_icon",
            "scripts/ — build_exe.py, build_icon.py, generate_true_prompt.py",
            "tests/ — pytest",
            "assets/ — keepcopy_icon.svg / .ico",
            "README.md, BUILD.md, FILTERS.md",
        ],
    )

    # --- 14. Тесты ---
    doc.add_heading("14. Тесты", level=1)
    doc.add_paragraph("pytest tests/ -v — без GUI, временные каталоги.")
    _add_table(
        doc,
        ["Файл", "Покрытие"],
        [
            ["test_backup.py", "режимы, ZIP, фильтры, max_size, disk full, multi-source"],
            ["test_autostart.py", "ярлык, sync, reconcile"],
            ["test_path_utils.py", "пути, Nuitka detection, nested sources"],
            ["test_task_dialog.py", "валидация, nested drop"],
            ["test_size_cache.py", "normalize кэша, calculate_size по источникам"],
            ["test_size_signal.py", "большие размеры через сигнал Qt"],
        ],
    )

    # --- 15. Критерии ---
    doc.add_heading("15. Критерии приёмки", level=1)
    for i, c in enumerate(
        [
            "GUI и все диалоги соответствуют §4; раскрытие источников и размеры — строго по §4.4.",
            "Три режима копирования, ZIP, фильтры, диск, errors/ — по §6.",
            "Расписание и очередь — по §7–8.",
            "settings.json с size_cache {total, sources} — по §9.",
            "Автозапуск + трей + правило закрытия — по §3 и §11.",
            "Сборка Nuitka в KeepCopy.dist — по §12.",
            "Все перечисленные pytest проходят.",
            "Документация README.md, BUILD.md, FILTERS.md на месте.",
        ],
        1,
    ):
        doc.add_paragraph(c, style="List Number")

    doc.add_page_break()
    doc.add_heading("Приложение А. Поток выполнения", level=1)
    flow = (
        "Запуск → settings.json → reconcile_autostart → MainWindow\n"
        "  (+ AppTray если --background или autostart)\n"
        "  → AutoScheduler(30s) + catch-up(5 мин при просрочке)\n"
        "  → scan_all_sizes\n"
        "Выполнить / планировщик → TaskQueue\n"
        "  → SizeScanWorker {total, sources} → BackupWorker → BackupEngine\n"
        "  → FileMatcher → disk check → copy → backup.log / errors/ → save tasks"
    )
    p = doc.add_paragraph()
    r = p.add_run(flow)
    r.font.name = "Consolas"
    r.font.size = Pt(9)

    doc.add_heading("Приложение Б. Чеклист «легко забыть»", level=1)
    _bullets(
        doc,
        [
            "Первый источник не в выпадающем списке; его размер при раскрытии — в строке задачи цветом detail.",
            "Свёрнуто — сумма обычным цветом.",
            "Нет анимации раскрытия строк.",
            "Фон кнопок действий = фон строки (select/hover/zebra).",
            "Вложенные источники в диалоге красные и не сохраняются.",
            "Ручной keep_changes копирует и «пропавшие» в назначении файлы.",
            "При autostart закрытие окна ≠ выход из процесса.",
            "Ярлык автозапуска с абсолютным путём — обновлять при каждом старте.",
        ],
    )

    return doc


def main() -> None:
    doc = build_document()
    doc.save(OUTPUT)
    print(f"Создан файл: {OUTPUT}")


if __name__ == "__main__":
    main()
